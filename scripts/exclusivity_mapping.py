#!/usr/bin/env python3

import os
import glob
from Bio import AlignIO
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX

base_dir = os.path.dirname(os.path.abspath(__file__))
alignment_files = glob.glob(f"{base_dir}/**/exclusivity_aligment_*.aln-clustalw", recursive=True)

def underline_run(run):
    rPr = run._element.get_or_add_rPr()
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

def wave_underline_run(run):
    rPr = run._element.get_or_add_rPr()
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "wave")
    rPr.append(u)

for alignment_path in alignment_files:
    folder = os.path.dirname(alignment_path)
    filename = os.path.basename(alignment_path)
    gene = filename.replace("exclusivity_aligment_", "").split("_S_pyogenes")[0]

    print(f"🔍 Processing exclusivity file for gene: {gene}")
    output_docx = os.path.join(folder, f"{gene}_exclusivity_mismatches_on_pyogenes.docx")

    try:
        alignment = AlignIO.read(alignment_path, "clustal")
    except Exception as e:
        print(f"❌ Failed to read {filename}: {e}")
        continue

    # Ensure correct identification of pyogenes vs target
    pyogenes_record = None
    exclusivity_record = None

    for record in alignment:
        if "pyogenes" in record.id.lower():
            pyogenes_record = record
        else:
            exclusivity_record = record

    if not pyogenes_record or not exclusivity_record:
        print(f"⚠️  Could not detect both S. pyogenes and exclusivity hit in: {filename}")
        continue

    print(f"   ↪ Reference (S. pyogenes): {pyogenes_record.id}")
    print(f"   ↪ Compared to (exclusivity): {exclusivity_record.id}")

    pyogenes_seq = str(pyogenes_record.seq)
    exclusivity_seq = str(exclusivity_record.seq)

    # === Build Word doc showing pyogenes with mismatches highlighted ===
    doc = Document()
    doc.add_heading(f"{gene}: S. pyogenes sequence with mismatches from {exclusivity_record.id}", level=1)
    para = doc.add_paragraph()

    for p_base, e_base in zip(pyogenes_seq, exclusivity_seq):
        run = para.add_run(p_base)
        if p_base == "-":
            wave_underline_run(run)
        elif p_base != e_base:
            underline_run(run)
            run.bold = True
            run.font.color.rgb = RGBColor(165, 42, 42)
            run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN

    doc.save(output_docx)
    print(f"📄 Output saved: {output_docx}\n")
