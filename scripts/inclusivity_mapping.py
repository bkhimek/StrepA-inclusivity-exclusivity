import os
import glob
import csv
from collections import defaultdict
from Bio import AlignIO
from Bio import SeqIO
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

# Use the script's own directory as base
base_dir = os.path.dirname(os.path.abspath(__file__))

# Find all alignments in this folder and subfolders
alignment_files = glob.glob(f"{base_dir}/**/*_alignment.aln", recursive=True)

for alignment_path in alignment_files:
    folder = os.path.dirname(alignment_path)
    basename = os.path.basename(alignment_path)
    gene = basename.split("_alignment")[0]

    print(f"🔍 Processing gene: {gene}")

    consensus_file = os.path.join(folder, f"{gene}_consensus.fasta")
    output_csv = os.path.join(folder, f"{gene}_SNP_summary_collapsed.csv")
    output_docx = os.path.join(folder, f"{gene}_consensus_highlighted.docx")

    # Skip if already processed
    if os.path.exists(output_csv) and os.path.exists(output_docx):
        print(f"⏭️  Skipping {gene}, already processed.")
        continue

    if not os.path.exists(consensus_file):
        print(f"⚠️  Consensus file not found: {consensus_file}")
        continue

    # === Load alignment and summarize SNPs ===
    alignment = AlignIO.read(alignment_path, "fasta")
    consensus_seq = str(alignment[0].seq)
    snp_summary = defaultdict(list)

    for record in alignment[1:]:
        strain_seq = str(record.seq)
        strain_id = record.id
        for pos, (cons_base, strain_base) in enumerate(zip(consensus_seq, strain_seq), start=1):
            if cons_base != strain_base:
                snp_summary[pos].append((strain_id, strain_base))

    # === Write SNP summary to CSV ===
    with open(output_csv, "w", newline="") as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow(["Position", "Consensus_Base", "Num_Differences", "Strain_IDs", "Strain_Bases"])
        for pos in sorted(snp_summary.keys()):
            strain_ids = ";".join([x[0] for x in snp_summary[pos]])
            strain_bases = ";".join([x[1] for x in snp_summary[pos]])
            writer.writerow([pos, consensus_seq[pos - 1], len(snp_summary[pos]), strain_ids, strain_bases])

    print(f"✅ SNP summary saved: {output_csv}")

    # === Create highlighted .docx ===
    doc = Document()
    para = doc.add_paragraph()

    for i, base in enumerate(consensus_seq, start=1):
        run = para.add_run(base)
        if i in snp_summary:
            count = len(snp_summary[i])
            if count == 1:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            elif count == 2:
                run.font.highlight_color = WD_COLOR_INDEX.PINK
            elif count > 2:
                run.font.highlight_color = WD_COLOR_INDEX.RED

    doc.save(output_docx)
    print(f"📄 Highlighted consensus saved: {output_docx}\n")

# === Auto-copy both .docx and .csv to Windows output folder ===
windows_output_dir = f"/mnt/c/Users/krist/OneDrive/Documents/Strep_A/alignments/{gene}/outputs"
os.makedirs(windows_output_dir, exist_ok=True)

try:
    import shutil
    shutil.copy(output_docx, os.path.join(windows_output_dir, os.path.basename(output_docx)))
    shutil.copy(output_csv, os.path.join(windows_output_dir, os.path.basename(output_csv)))
    print(f"📁 Copied files to Windows: {windows_output_dir}\n")
except Exception as e:
    print(f"⚠️  Failed to copy to Windows folder: {e}")
